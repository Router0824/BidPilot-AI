import os
import hashlib
import json
import shutil
import aiofiles
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from app.domain.models import Document, DocumentPage, ParseStatus
from app.core.config import settings


class DocumentService:
    def __init__(self, db: AsyncSession):
        self.db = db

    async def list_documents(self, project_id: str) -> list[Document]:
        result = await self.db.execute(
            select(Document).where(Document.project_id == project_id).order_by(Document.created_at.desc())
        )
        return list(result.scalars().all())

    async def get_document(self, document_id: str) -> Document | None:
        result = await self.db.execute(select(Document).where(Document.id == document_id))
        return result.scalar_one_or_none()

    async def upload_document(self, project_id: str, filename: str, content: bytes, doc_type: str, user: dict) -> Document:
        file_hash = hashlib.sha256(content).hexdigest()
        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
        file_dir = os.path.join(settings.UPLOAD_DIR, project_id)
        os.makedirs(file_dir, exist_ok=True)
        file_path = os.path.join(file_dir, os.path.basename(filename))
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(content)

        doc = Document(
            project_id=project_id,
            name=filename,
            document_type=doc_type,
            file_path=file_path,
            file_size=len(content),
            file_hash=file_hash,
            uploaded_by=user["id"],
        )
        self.db.add(doc)
        await self.db.flush()
        return doc

    async def create_upload_session(
        self,
        project_id: str,
        filename: str,
        total_size: int,
        file_hash: str,
        doc_type: str,
        user: dict,
        chunk_size: int = 8 * 1024 * 1024,
    ) -> dict:
        import uuid

        session_id = str(uuid.uuid4())
        session_dir = os.path.join(settings.UPLOAD_DIR, project_id, ".upload_sessions", session_id)
        os.makedirs(session_dir, exist_ok=True)
        meta = {
            "id": session_id,
            "project_id": project_id,
            "filename": os.path.basename(filename),
            "total_size": total_size,
            "file_hash": file_hash,
            "document_type": doc_type,
            "uploaded_by": user["id"],
            "chunk_size": chunk_size,
            "received_chunks": [],
        }
        await self._write_upload_meta(session_dir, meta)
        return meta

    async def upload_chunk(self, project_id: str, session_id: str, chunk_index: int, content: bytes) -> dict:
        session_dir = os.path.join(settings.UPLOAD_DIR, project_id, ".upload_sessions", session_id)
        meta = await self._read_upload_meta(session_dir)
        chunk_path = os.path.join(session_dir, f"{chunk_index:08d}.part")
        async with aiofiles.open(chunk_path, "wb") as f:
            await f.write(content)
        received = set(meta.get("received_chunks") or [])
        received.add(chunk_index)
        meta["received_chunks"] = sorted(received)
        await self._write_upload_meta(session_dir, meta)
        return {"upload_session_id": session_id, "chunk_index": chunk_index, "received_chunks": meta["received_chunks"]}

    async def complete_upload_session(self, project_id: str, session_id: str, user: dict) -> Document:
        session_dir = os.path.join(settings.UPLOAD_DIR, project_id, ".upload_sessions", session_id)
        meta = await self._read_upload_meta(session_dir)
        final_dir = os.path.join(settings.UPLOAD_DIR, project_id)
        os.makedirs(final_dir, exist_ok=True)
        final_path = os.path.join(final_dir, meta["filename"])
        hasher = hashlib.sha256()
        total_written = 0

        part_files = sorted(p for p in os.listdir(session_dir) if p.endswith(".part"))
        async with aiofiles.open(final_path, "wb") as out:
            for part_name in part_files:
                part_path = os.path.join(session_dir, part_name)
                async with aiofiles.open(part_path, "rb") as part:
                    while True:
                        chunk = await part.read(1024 * 1024)
                        if not chunk:
                            break
                        hasher.update(chunk)
                        total_written += len(chunk)
                        await out.write(chunk)

        actual_hash = hasher.hexdigest()
        if meta.get("total_size") and total_written != int(meta["total_size"]):
            os.remove(final_path)
            raise ValueError(f"文件大小校验失败：期望 {meta['total_size']}，实际 {total_written}")
        if meta.get("file_hash") and actual_hash.lower() != str(meta["file_hash"]).lower():
            os.remove(final_path)
            raise ValueError("SHA-256 校验失败")

        doc = Document(
            project_id=project_id,
            name=meta["filename"],
            document_type=meta["document_type"],
            file_path=final_path,
            file_size=total_written,
            file_hash=actual_hash,
            uploaded_by=user["id"],
        )
        self.db.add(doc)
        shutil.rmtree(session_dir, ignore_errors=True)
        await self.db.flush()
        return doc

    async def _read_upload_meta(self, session_dir: str) -> dict:
        meta_path = os.path.join(session_dir, "meta.json")
        if not os.path.exists(meta_path):
            raise ValueError("上传会话不存在")
        async with aiofiles.open(meta_path, "r", encoding="utf-8") as f:
            return json.loads(await f.read())

    async def _write_upload_meta(self, session_dir: str, meta: dict) -> None:
        async with aiofiles.open(os.path.join(session_dir, "meta.json"), "w", encoding="utf-8") as f:
            await f.write(json.dumps(meta, ensure_ascii=False))

    async def parse_document(self, document_id: str) -> Document:
        doc = await self.get_document(document_id)
        if not doc:
            raise ValueError("Document not found")
        doc.parse_status = ParseStatus.PARSING.value
        await self.db.flush()

        try:
            pages = await self._do_parse(doc)
            for p in pages:
                self.db.add(p)
            doc.parse_status = ParseStatus.COMPLETED.value
            doc.page_count = len(pages)
        except Exception:
            doc.parse_status = ParseStatus.FAILED.value
        await self.db.flush()
        return doc

    async def _do_parse(self, doc: Document) -> list[DocumentPage]:
        pages = []
        ext = doc.name.lower().split(".")[-1] if "." in doc.name else ""

        if ext == "pdf":
            pages = await self._parse_pdf(doc)
        elif ext in ("docx", "doc"):
            pages = await self._parse_docx(doc)
        elif ext in ("xlsx", "xls"):
            pages = await self._parse_xlsx(doc)
        elif ext in ("txt", "md"):
            pages = await self._parse_text(doc)
        else:
            pages = await self._parse_text(doc)
        return pages

    async def _parse_pdf(self, doc: Document) -> list[DocumentPage]:
        import pdfplumber
        pages = []
        with pdfplumber.open(doc.file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                tables = page.extract_tables() or []
                dp = DocumentPage(
                    document_id=doc.id,
                    page_number=i + 1,
                    text=text,
                    parse_method="native",
                    table_count=len(tables),
                    quality_score=0.9,
                )
                pages.append(dp)
        return pages

    async def _parse_docx(self, doc: Document) -> list[DocumentPage]:
        from docx import Document as DocxDocument  # noqa: N811
        d = DocxDocument(doc.file_path)
        full_text = "\n".join(p.text for p in d.paragraphs if p.text)
        dp = DocumentPage(
            document_id=doc.id, page_number=1,
            text=full_text, parse_method="ooxml", table_count=len(d.tables),
            quality_score=0.95,
        )
        return [dp]

    async def _parse_xlsx(self, doc: Document) -> list[DocumentPage]:
        from openpyxl import load_workbook
        wb = load_workbook(doc.file_path, data_only=True)
        all_text = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            all_text.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(c) if c is not None else "" for c in row)
                all_text.append(row_text)
        dp = DocumentPage(
            document_id=doc.id, page_number=1,
            text="\n".join(all_text), parse_method="ooxml",
            table_count=len(wb.sheetnames), quality_score=0.9,
        )
        return [dp]

    async def _parse_text(self, doc: Document) -> list[DocumentPage]:
        async with aiofiles.open(doc.file_path, "r", encoding="utf-8", errors="replace") as f:
            text = await f.read()
        dp = DocumentPage(
            document_id=doc.id, page_number=1,
            text=text, parse_method="native", quality_score=1.0,
        )
        return [dp]

    async def get_pages(self, document_id: str) -> list[DocumentPage]:
        result = await self.db.execute(
            select(DocumentPage).where(DocumentPage.document_id == document_id).order_by(DocumentPage.page_number)
        )
        return list(result.scalars().all())

    async def delete_document(self, document_id: str) -> bool:
        doc = await self.get_document(document_id)
        if not doc:
            return False
        if doc.file_path and os.path.exists(doc.file_path):
            os.remove(doc.file_path)
        await self.db.delete(doc)
        await self.db.flush()
        return True
