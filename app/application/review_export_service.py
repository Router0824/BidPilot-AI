import re

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func, and_
from app.domain.models import (
    ReviewRun, ReviewFinding, Requirement, DraftVersion, OutlineSection,
    ProjectFact, RequirementStatus, RiskLevel, ConfirmStatus, FixAttempt
)
from app.schemas import ExportRequest


class ReviewService:
    def __init__(self, db: AsyncSession):
        self.db = db

    async def list_reviews(self, project_id: str) -> list[ReviewRun]:
        result = await self.db.execute(
            select(ReviewRun).where(ReviewRun.project_id == project_id).order_by(ReviewRun.created_at.desc())
        )
        return list(result.scalars().all())

    async def run_review(self, project_id: str, review_type: str = "full") -> ReviewRun:
        review_run = ReviewRun(project_id=project_id, review_type=review_type, status="running")
        self.db.add(review_run)
        await self.db.flush()

        findings = []

        if review_type in ("full", "coverage"):
            findings += await self._check_coverage(project_id, review_run.id)
        if review_type in ("full", "consistency"):
            findings += await self._check_consistency(project_id, review_run.id)
        if review_type in ("full", "citation"):
            findings += await self._check_citations(project_id, review_run.id)
        if review_type in ("full", "overcommit"):
            findings += await self._check_overcommit(project_id, review_run.id)

        review_run.findings = [f.id for f in findings]
        review_run.status = "completed"
        await self.db.flush()
        return review_run

    async def _check_coverage(self, project_id: str, review_run_id: str) -> list[ReviewFinding]:
        findings = []
        result = await self.db.execute(
            select(Requirement).where(
                and_(Requirement.project_id == project_id, Requirement.status != RequirementStatus.RESPONDED.value)
            )
        )
        uncovered = result.scalars().all()
        for req in uncovered:
            f = ReviewFinding(
                review_run_id=review_run_id,
                finding_type="missing_requirement",
                risk_level=req.risk_level,
                description=f"要求未响应：{req.requirement_text[:200]}",
                location=f"requirement:{req.id}",
                suggestion="请确认是否有意在标书中省略此要求，或为要求分配响应章节",
            )
            self.db.add(f)
            findings.append(f)
        return findings

    async def _check_consistency(self, project_id: str, review_run_id: str) -> list[ReviewFinding]:
        findings = []
        facts_result = await self.db.execute(
            select(ProjectFact).where(ProjectFact.project_id == project_id)
        )
        facts = {f.fact_key: f.fact_value for f in facts_result.scalars().all()}

        sections_result = await self.db.execute(
            select(OutlineSection).where(OutlineSection.project_id == project_id)
        )
        sections = sections_result.scalars().all()

        for section in sections:
            if section.current_version_id:
                draft = (await self.db.execute(
                    select(DraftVersion).where(DraftVersion.id == section.current_version_id)
                )).scalar_one_or_none()
                if draft and draft.content:
                    for key, val in facts.items():
                        if val and val not in draft.content:
                            f = ReviewFinding(
                                review_run_id=review_run_id,
                                finding_type="numeric_inconsistency" if re.search(r"\d", str(val)) else "internal_conflict",
                                risk_level="medium",
                                description=f"章节「{section.title}」中未发现项目事实「{key}」({val})",
                                location=f"section:{section.id}",
                                suggestion=f"请确认章节中是否包含 {key} 相关信息",
                            )
                            self.db.add(f)
                            findings.append(f)
        return findings

    async def _check_citations(self, project_id: str, review_run_id: str) -> list[ReviewFinding]:
        findings = []
        sections_result = await self.db.execute(
            select(OutlineSection).where(
                and_(OutlineSection.project_id == project_id, OutlineSection.current_version_id.isnot(None))
            )
        )
        for section in sections_result.scalars().all():
            draft = (await self.db.execute(
                select(DraftVersion).where(DraftVersion.id == section.current_version_id)
            )).scalar_one_or_none()
            if draft and draft.citations:
                for cite in draft.citations:
                    if cite.get("status") == "unverified":
                        f = ReviewFinding(
                            review_run_id=review_run_id,
                            finding_type="citation_missing",
                            risk_level="medium",
                            description=f"章节「{section.title}」中存在未验证引用：{cite.get('source', '')}",
                            location=f"section:{section.id}",
                            suggestion="请确认引用来源的有效性",
                        )
                        self.db.add(f)
                        findings.append(f)
        return findings

    async def _check_overcommit(self, project_id: str, review_run_id: str) -> list[ReviewFinding]:
        findings = []
        overcommit_keywords = ["准确率", "99", "100%", "零故障", "绝对安全", "万无一失", "保证中标"]
        sections_result = await self.db.execute(
            select(OutlineSection).where(
                and_(OutlineSection.project_id == project_id, OutlineSection.current_version_id.isnot(None))
            )
        )
        for section in sections_result.scalars().all():
            draft = (await self.db.execute(
                select(DraftVersion).where(DraftVersion.id == section.current_version_id)
            )).scalar_one_or_none()
            if draft and draft.content:
                for kw in overcommit_keywords:
                    if kw in draft.content:
                        f = ReviewFinding(
                            review_run_id=review_run_id,
                            finding_type="unsupported_claim",
                            risk_level="high",
                            description=f"章节「{section.title}」中存在过度承诺嫌疑：包含「{kw}」",
                            location=f"section:{section.id}",
                            suggestion="请确认该表述是否有企业材料支撑，若无请删除或修改",
                        )
                        self.db.add(f)
                        findings.append(f)
                        break
        return findings

    async def list_findings(self, review_run_id: str) -> list[ReviewFinding]:
        result = await self.db.execute(
            select(ReviewFinding).where(ReviewFinding.review_run_id == review_run_id)
        )
        return list(result.scalars().all())

    async def update_finding(self, finding_id: str, status: str, ignore_reason: str | None = None) -> ReviewFinding | None:
        result = await self.db.execute(select(ReviewFinding).where(ReviewFinding.id == finding_id))
        finding = result.scalar_one_or_none()
        if not finding:
            return None
        finding.status = status
        if ignore_reason:
            finding.ignore_reason = ignore_reason
        await self.db.flush()
        return finding


HIGH_RISK_ISSUE_TYPES = {
    "qualification_risk",
    "schedule_conflict",
    "addendum_conflict",
    "unsupported_claim",
    "numeric_inconsistency",
}

AUTO_FIXABLE_TYPES = {
    "citation_missing",
}


class FixerService:
    MAX_ATTEMPTS_PER_SECTION = 2

    def __init__(self, db: AsyncSession):
        self.db = db

    async def list_attempts(self, project_id: str) -> list[FixAttempt]:
        result = await self.db.execute(
            select(FixAttempt).where(FixAttempt.project_id == project_id).order_by(FixAttempt.created_at.desc())
        )
        return list(result.scalars().all())

    async def fix_issue(
        self,
        project_id: str,
        issue_id: str,
        user: dict,
        apply: bool = True,
    ) -> FixAttempt:
        issue = (await self.db.execute(select(ReviewFinding).where(ReviewFinding.id == issue_id))).scalar_one_or_none()
        if not issue:
            raise ValueError("审查问题不存在")

        section = await self._section_for_issue(project_id, issue)
        draft = await self._draft_for_section(section)
        issue_type = issue.finding_type or "internal_conflict"
        risk_level = issue.risk_level or "medium"
        auto_fix_allowed = self._auto_fix_allowed(issue_type, risk_level)
        requires_confirmation = not auto_fix_allowed
        before = draft.content if draft and draft.content else ""

        section_id = section.id if section else None
        attempt_no = await self._next_attempt_no(project_id, section_id)
        if section_id and attempt_no > self.MAX_ATTEMPTS_PER_SECTION:
            return await self._record_attempt(
                project_id=project_id,
                issue=issue,
                section_id=section_id,
                draft=draft,
                before=before,
                after=before,
                reason="该章节自动修正次数已达到上限，停止循环并转人工处理",
                status="manual_required",
                attempt_no=attempt_no,
                auto_fix_allowed=False,
                requires_confirmation=True,
            )

        if not auto_fix_allowed:
            return await self._record_attempt(
                project_id=project_id,
                issue=issue,
                section_id=section_id,
                draft=draft,
                before=before,
                after=before,
                reason="高风险或内容承诺类问题必须人工确认，Fixer 不自动改写",
                status="manual_required",
                attempt_no=attempt_no,
                auto_fix_allowed=False,
                requires_confirmation=True,
            )

        after = self._apply_safe_fix(issue, section, before)
        status = "applied" if apply and draft and after != before else "suggested"
        attempt = await self._record_attempt(
            project_id=project_id,
            issue=issue,
            section_id=section_id,
            draft=draft,
            before=before,
            after=after,
            reason="低风险引用/格式问题，追加待确认引用说明，不编造资格、报价、工期或法律承诺",
            status=status,
            attempt_no=attempt_no,
            auto_fix_allowed=True,
            requires_confirmation=False,
        )
        if apply and draft and after != before:
            import datetime
            draft.content = after
            draft.word_count = len(after)
            attempt.applied_at = datetime.datetime.now(datetime.timezone.utc)
            issue.status = "resolved"
            await self.db.flush()
        return attempt

    async def _section_for_issue(self, project_id: str, issue: ReviewFinding) -> OutlineSection | None:
        location = issue.location or ""
        if location.startswith("section:"):
            section_id = location.split(":", 1)[1]
            return (await self.db.execute(select(OutlineSection).where(OutlineSection.id == section_id))).scalar_one_or_none()
        if location.startswith("requirement:"):
            req_id = location.split(":", 1)[1]
            req = (await self.db.execute(select(Requirement).where(Requirement.id == req_id))).scalar_one_or_none()
            if req and req.response_section_id:
                return (await self.db.execute(select(OutlineSection).where(OutlineSection.id == req.response_section_id))).scalar_one_or_none()
            title_hint = "公司资质" if req and req.requirement_type == "qualification" else None
            if title_hint:
                section = (await self.db.execute(
                    select(OutlineSection).where(
                        OutlineSection.project_id == project_id,
                        OutlineSection.title.contains(title_hint),
                    ).limit(1)
                )).scalar_one_or_none()
                if section:
                    return section
        return (await self.db.execute(
            select(OutlineSection).where(
                OutlineSection.project_id == project_id,
                OutlineSection.current_version_id.isnot(None),
            ).order_by(OutlineSection.sort_order).limit(1)
        )).scalar_one_or_none()

    async def _draft_for_section(self, section: OutlineSection | None) -> DraftVersion | None:
        if not section or not section.current_version_id:
            return None
        return (await self.db.execute(
            select(DraftVersion).where(DraftVersion.id == section.current_version_id)
        )).scalar_one_or_none()

    async def _next_attempt_no(self, project_id: str, section_id: str | None) -> int:
        if not section_id:
            return 1
        result = await self.db.execute(
            select(FixAttempt).where(
                FixAttempt.project_id == project_id,
                FixAttempt.section_id == section_id,
                FixAttempt.status == "applied",
            )
        )
        return len(result.scalars().all()) + 1

    def _auto_fix_allowed(self, issue_type: str, risk_level: str) -> bool:
        if risk_level == "high" or issue_type in HIGH_RISK_ISSUE_TYPES:
            return False
        return issue_type in AUTO_FIXABLE_TYPES

    def _apply_safe_fix(self, issue: ReviewFinding, section: OutlineSection | None, before: str) -> str:
        note = (
            "\n\n> 引用补充：本段引用来源需由编制人员核对原始材料后确认；"
            "当前自动修正仅增加引用待确认标记，不新增事实承诺。"
        )
        if note.strip() in before:
            return before
        if before.strip():
            return before.rstrip() + note
        section_title = section.title if section else "待绑定章节"
        return f"## {section_title}\n\n【待确认】请补充本章节内容及引用来源。{note}"

    async def _record_attempt(
        self,
        project_id: str,
        issue: ReviewFinding,
        section_id: str | None,
        draft: DraftVersion | None,
        before: str,
        after: str,
        reason: str,
        status: str,
        attempt_no: int,
        auto_fix_allowed: bool,
        requires_confirmation: bool,
    ) -> FixAttempt:
        attempt = FixAttempt(
            project_id=project_id,
            issue_id=issue.id,
            section_id=section_id,
            draft_version_id=draft.id if draft else None,
            issue_type=issue.finding_type,
            risk_level=issue.risk_level,
            auto_fix_allowed=auto_fix_allowed,
            status=status,
            reason=reason,
            before_content=before,
            after_content=after,
            diff=self._simple_diff(before, after),
            attempt_no=attempt_no,
            requires_human_confirmation=requires_confirmation,
        )
        self.db.add(attempt)
        await self.db.flush()
        return attempt

    def _simple_diff(self, before: str, after: str) -> dict:
        before_lines = before.splitlines()
        after_lines = after.splitlines()
        added = [line for line in after_lines if line not in before_lines]
        removed = [line for line in before_lines if line not in after_lines]
        return {"added": added, "removed": removed}


class ExportService:
    def __init__(self, db: AsyncSession):
        self.db = db

    async def export(self, project_id: str, export_type: str, fmt: str) -> str:
        import os
        export_dir = os.path.join("uploads", project_id, "exports")
        os.makedirs(export_dir, exist_ok=True)

        if export_type == "requirements":
            return await self._export_requirements(project_id, export_dir, fmt)
        elif export_type == "outline":
            return await self._export_outline(project_id, export_dir, fmt)
        elif export_type == "full_document":
            return await self._export_full(project_id, export_dir, fmt)
        elif export_type == "risk_list":
            return await self._export_risk_list(project_id, export_dir, fmt)
        else:
            return await self._export_draft(project_id, export_dir, fmt)

    async def _export_requirements(self, project_id: str, export_dir: str, fmt: str) -> str:
        result = await self.db.execute(
            select(Requirement).where(Requirement.project_id == project_id).order_by(Requirement.risk_level.desc())
        )
        reqs = result.scalars().all()

        if fmt == "docx":
            filepath = f"{export_dir}/requirements.docx"
            self._write_requirements_docx(filepath, reqs)
            return filepath
        if fmt == "markdown":
            lines = ["# 要求与风险矩阵\n"]
            lines.append("| ID | 类型 | 风险 | 硬性 | 要求内容 | 状态 |")
            lines.append("|----|------|------|------|----------|------|")
            for r in reqs:
                mandatory = "是" if r.mandatory else "否"
                lines.append(f"| {r.id[:8]} | {r.requirement_type} | {r.risk_level} | {mandatory} | {r.requirement_text[:100]} | {r.status} |")
            content = "\n".join(lines)
        elif fmt == "xlsx":
            content = await self._export_xlsx(reqs, ["ID", "类型", "风险", "硬性", "要求内容", "状态"],
                lambda r: [r.id[:8], r.requirement_type, r.risk_level, "是" if r.mandatory else "否", r.requirement_text[:100], r.status])
        else:
            content = "\n".join(f"{r.requirement_type}\t{r.risk_level}\t{r.requirement_text}" for r in reqs)

        filepath = f"{export_dir}/requirements.{fmt}"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        return filepath

    async def _export_outline(self, project_id: str, export_dir: str, fmt: str) -> str:
        result = await self.db.execute(
            select(OutlineSection).where(OutlineSection.project_id == project_id).order_by(OutlineSection.sort_order)
        )
        sections = result.scalars().all()
        if fmt == "docx":
            filepath = f"{export_dir}/outline.docx"
            self._write_outline_docx(filepath, sections)
            return filepath
        lines = ["# 技术标大纲\n"]
        for s in sections:
            indent = "  " * (s.level - 1)
            lines.append(f"{indent}- {s.title}")
        content = "\n".join(lines)
        filepath = f"{export_dir}/outline.{fmt}"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        return filepath

    async def _export_full(self, project_id: str, export_dir: str, fmt: str) -> str:
        result = await self.db.execute(
            select(OutlineSection).where(OutlineSection.project_id == project_id).order_by(OutlineSection.sort_order)
        )
        sections = result.scalars().all()
        if fmt == "docx":
            filepath = f"{export_dir}/full_document.docx"
            await self._write_full_docx(filepath, sections)
            return filepath
        lines = ["# 技术标书\n"]
        for s in sections:
            prefix = "#" * min(s.level, 6)
            lines.append(f"\n{prefix} {s.title}\n")
            if s.current_version_id:
                draft = (await self.db.execute(
                    select(DraftVersion).where(DraftVersion.id == s.current_version_id)
                )).scalar_one_or_none()
                if draft and draft.content:
                    lines.append(draft.content)
                    lines.append("")
        content = "\n".join(lines)
        filepath = f"{export_dir}/full_document.{fmt}"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        return filepath

    async def _export_risk_list(self, project_id: str, export_dir: str, fmt: str) -> str:
        result = await self.db.execute(
            select(Requirement).where(
                and_(Requirement.project_id == project_id, Requirement.risk_level.in_(["high", "medium"]))
            ).order_by(Requirement.risk_level.desc())
        )
        reqs = result.scalars().all()
        if fmt == "docx":
            filepath = f"{export_dir}/risk_list.docx"
            self._write_risk_list_docx(filepath, reqs)
            return filepath
        lines = ["# 风险清单\n"]
        for r in reqs:
            lines.append(f"## [{r.risk_level.upper()}] {r.requirement_text[:200]}")
            lines.append(f"- 类型：{r.requirement_type}")
            lines.append(f"- 来源：{r.source_document_id} 第{r.source_page}页")
            lines.append(f"- 状态：{r.status}")
            lines.append("")
        content = "\n".join(lines)
        filepath = f"{export_dir}/risk_list.{fmt}"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        return filepath

    async def _export_draft(self, project_id: str, export_dir: str, fmt: str) -> str:
        return await self._export_full(project_id, export_dir, fmt)

    async def _export_xlsx(self, items, headers, row_fn) -> str:
        rows = ["\t".join(headers)]
        rows.extend("\t".join(str(c) for c in row_fn(i)) for i in items)
        return "\n".join(rows)

    def _new_docx(self):
        from docx import Document as DocxDocument  # noqa: N811
        from docx.oxml.ns import qn

        doc = DocxDocument()
        styles = doc.styles
        for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
            style = styles[style_name]
            style.font.name = "Noto Serif CJK SC"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif CJK SC")
        return doc

    def _add_markdownish_paragraphs(self, doc, content: str) -> None:
        from docx.enum.text import WD_COLOR_INDEX

        for line in (content or "").splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            else:
                paragraph = doc.add_paragraph()
                cursor = 0
                for match in re.finditer(r"【.+?】", stripped):
                    if match.start() > cursor:
                        paragraph.add_run(stripped[cursor:match.start()])
                    run = paragraph.add_run(match.group(0))
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    cursor = match.end()
                if cursor < len(stripped):
                    paragraph.add_run(stripped[cursor:])

    def _write_requirements_docx(self, filepath: str, reqs: list[Requirement]) -> None:
        doc = self._new_docx()
        doc.add_heading("要求与风险矩阵", level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["ID", "类型", "风险", "硬性", "要求内容", "状态"]
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header
        for req in reqs:
            cells = table.add_row().cells
            values = [req.id[:8], req.requirement_type, req.risk_level, "是" if req.mandatory else "否", req.requirement_text[:300], req.status]
            for idx, value in enumerate(values):
                cells[idx].text = str(value)
        doc.save(filepath)

    def _write_outline_docx(self, filepath: str, sections: list[OutlineSection]) -> None:
        doc = self._new_docx()
        doc.add_heading("技术标大纲", level=1)
        for section in sections:
            doc.add_paragraph(f"{'  ' * max(0, section.level - 1)}{section.sort_order}. {section.title}")
        doc.save(filepath)

    async def _write_full_docx(self, filepath: str, sections: list[OutlineSection]) -> None:
        doc = self._new_docx()
        doc.add_heading("技术标书", level=1)
        for section in sections:
            doc.add_heading(section.title, level=min(max(section.level, 1), 3))
            if section.current_version_id:
                draft = (await self.db.execute(
                    select(DraftVersion).where(DraftVersion.id == section.current_version_id)
                )).scalar_one_or_none()
                if draft and draft.content:
                    self._add_markdownish_paragraphs(doc, draft.content)
                if draft and draft.citations:
                    doc.add_paragraph("引用来源：")
                    for cite in draft.citations:
                        doc.add_paragraph(
                            f"{cite.get('source', '')} 第{cite.get('page', '-')}页：{cite.get('snippet', '')}",
                            style="List Bullet",
                        )
        doc.save(filepath)

    def _write_risk_list_docx(self, filepath: str, reqs: list[Requirement]) -> None:
        doc = self._new_docx()
        doc.add_heading("风险清单", level=1)
        for req in reqs:
            doc.add_heading(f"[{req.risk_level.upper()}] {req.requirement_text[:80]}", level=2)
            doc.add_paragraph(f"类型：{req.requirement_type}")
            doc.add_paragraph(f"来源：{req.source_document_id or '-'} 第{req.source_page or '-'}页")
            doc.add_paragraph(f"状态：{req.status}")
        doc.save(filepath)
